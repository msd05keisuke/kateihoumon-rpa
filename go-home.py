import os
import time
import datetime
# Seleniumライブラリ
from selenium import webdriver
# Selenium待機ライブラリ
from selenium.webdriver.support import expected_conditions
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
# Excel操作ライブラリ
import xlwings as xw
# Word操作ライブラリ
from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
# 画像加工ライブラリ
from PIL import Image


# webドライバ(chrome)
browser = webdriver.Chrome('/usr/local/bin/chromedriver')
browser.maximize_window()  # 画面サイズ:最大

# タイムアウト値を指定
wait = WebDriverWait(browser, 10)  # 10秒

# GoogleMapを開く
browser.get('https://www.google.co.jp/maps/@35.710067,139.8085117,17z?hl=ja')

# Excelシートを開く
ex_path = 'go_home.xlsx'
ex_b = xw.Book(ex_path)
ex_s = ex_b.sheets[0]

# 学籍番号が存在している間、処理を繰り返す
i = 2
while ex_s.cells(i, 1).value is not None:

    num = str(ex_s.cells(i, 1).value)  # 学籍番号
    student = str(ex_s.cells(i, 2).value)  # 生徒氏名
    parent = str(ex_s.cells(i, 3).value)  # 保護者氏名
    address = str(ex_s.cells(i, 4).value)  # 住所
    tel = str(ex_s.cells(i, 5).value)  # 電話番号
    day = str(ex_s.cells(i, 6).value)  # 訪問希望日
    flg = True  # タイムアウトチェックフラグ

    # word関連
    document = Document()  # 新規作成
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run(str(datetime.date.today()))  # 今日の日付

    heading = document.add_heading('家庭訪問計画書', 0)  # タイトル追加
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    records = (
        ('学籍番号', num),
        ('生徒氏名', student),
        ('保護者', parent),
        ('住所', address),
        ('電話番号', tel),
        ('訪問希望日', day),

    )

    table = document.add_table(rows=0, cols=2, style='Table Grid')  # 表の作成
    for key, value in records:
        row_cells = table.add_row().cells
        row_cells[0].width = Mm(30)
        row_cells[1].width = Mm(155)
        row_cells[0].text = key
        row_cells[1].text = value

    # 地図の取得
    if browser.find_element_by_xpath('//span[@id="sb_cb50"]').is_displayed():
        browser.find_element_by_xpath('//span[@id="sb_cb50"]').click()  # ×ボタンが出ていたらクリック
    element = browser.find_element_by_xpath('//input[@id="searchboxinput"]')  # 検索ボックス
    element.clear()  # 検索ボックスの要素をクリアする
    element.send_keys(address)  # 検索ボックスへ住所を追加
    browser.find_element_by_xpath('//button[@id="searchbox-searchbutton"]').click()  # 検索ボタンをクリック

    # 住所が存在しなかった場合(タイムアウト)の処理
    try:
        # 住所が取得されるまで待機
        wait.until(expected_conditions.visibility_of_element_located(
            (By.XPATH, '//*[@id="pane"]/div/div[1]/div/div/div[2]/div[1]/div[1]/div[1]/h1')))
        time.sleep(3)

    except:
        flg = False

    if flg:
        # 画面キャプチャ
        browser.save_screenshot('img/tmp.png')

        # 画像のトリミング
        im1 = Image.open('img/tmp.png')
        im2 = im1.crop((520, 100, 1670, 895)).save('img/tmp2.png', quality=90)  # 左, 上, 右, 下

        # 地図画像の追加
        heading = document.add_heading('【周辺地図】', level=3)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_picture('img/tmp2.png', width=Mm(185))  # 幅185ミリ

        filename = num + '.docx'

        # 成功をExcelシートに記入
        ex_s.cells(i, 7).value = 'OK'
    else:
        filename = 'エラー' + num + '.docx'
        # 失敗ををExcelシートに記入
        ex_s.cells(i, 7).value = 'error'

    i += 1
    document.save(os.path.join('word', filename))



browser.close()



